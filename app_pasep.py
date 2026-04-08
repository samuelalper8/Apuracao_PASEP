"""
ConPrev — Super App de Auditoria do PASEP (Fisco-Federal) - Versão Ouro
Ofício Executivo: Tabela DRE, Filtro de Dízimas (Zeros Limpos) e Cloud Ready
"""
import streamlit as st
import pandas as pd
import re
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# ── Configurações e Sessão ───────────────────────────────────────────────────
st.set_page_config(page_title="ConPrev Premium CND", page_icon="⚖️", layout="wide")

if 'df_clientes' not in st.session_state:
    st.session_state.df_clientes = pd.DataFrame()
if 'df_representantes' not in st.session_state:
    st.session_state.df_representantes = pd.DataFrame()
if 'num_convenios' not in st.session_state:
    st.session_state.num_convenios = 1

# Dicionário para tradução automática do mês
MESES_MAP = {
    '01': 'Janeiro', '02': 'Fevereiro', '03': 'Março', '04': 'Abril',
    '05': 'Maio', '06': 'Junho', '07': 'Julho', '08': 'Agosto',
    '09': 'Setembro', '10': 'Outubro', '11': 'Novembro', '12': 'Dezembro'
}

# ── Motores de Formatação e Entrada ───────────────────────────────────────────
def formatar_moeda(valor, force_deduction_sign=False):
    if valor is None: return "R$ 0,00"
    
    # BLINDAGEM CONTÁBIL: Extermina lixos de ponto flutuante (ex: -0.000000001)
    valor = round(float(valor), 2)
    
    if valor == 0.0:
        return "R$ 0,00" # Garante que zero absoluto nunca terá sinal ou cor
        
    is_negativo = valor < 0 or (force_deduction_sign and valor > 0)
    abs_val = abs(valor)
    texto = f"R$ {abs_val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"-{texto}" if is_negativo else texto

def extrair_numero(texto):
    if not texto: return 0.0
    return float(texto.strip().replace('.', '').replace(',', '.').replace('_', '.'))

def entrada_financeira(label, valor_default=0.0, key_suffix="", label_visibility="visible"):
    str_default = f"{valor_default:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") if valor_default else "0,00"
    texto_digitado = st.text_input(label, value=str_default, key=f"{label}_{key_suffix}", label_visibility=label_visibility)
    numeros = re.sub(r'[^\d-]', '', str(texto_digitado))
    if not numeros or numeros == '-': return 0.0
    return float(numeros) / 100.0

def obter_data_formatada():
    hoje = datetime.now()
    mes_str = MESES_MAP[f"{hoje.month:02d}"].lower()
    return f"Goiânia - GO, {hoje.day:02d} de {mes_str} de {hoje.year}."

def processar_lote_csvs(arquivos_upload):
    dados_lote = []
    for arquivo in arquivos_upload:
        texto = arquivo.getvalue().decode('latin1', errors='ignore')
        nome_arquivo = arquivo.name
        
        municipio_nome = nome_arquivo.replace('.csv', '')
        uf_extraida = "GO" # Default
        match = re.search(r'(GO|MS|TO) - (\d{4}-\d{2}) - (.*)', municipio_nome)
        if match:
            uf_extraida = match.group(1).strip()
            municipio_nome = f"{match.group(3).strip()}"
            
        cliente_data = {
            'Arquivo': nome_arquivo, 'Cliente': municipio_nome, 'UF': uf_extraida, 'MesRef': '',
            'FPM': 0.0, 'ITR': 0.0, 'FEP': 0.0, 'LC_176': 0.0, 'Simples_Nacional': 0.0, 'CFM': 0.0, 'CIDE': 0.0
        }
        
        secao_atual = None
        dentro_totais = False
        for linha in texto.split('\n'):
            linha_upper = linha.upper()
            
            if not cliente_data['MesRef']:
                match_data = re.search(r'\b\d{2}\.(\d{2})\.(\d{4})\b', linha_upper)
                if match_data:
                    mes = MESES_MAP.get(match_data.group(1), '')
                    ano = match_data.group(2)
                    if mes and ano:
                        cliente_data['MesRef'] = f"{mes}/{ano}"
            
            if '- FUNDO DE PARTICIPACAO' in linha_upper: secao_atual = 'FPM'
            elif '- IMPOSTO TERRITORIAL RURAL' in linha_upper: secao_atual = 'ITR'
            elif '- FUNDO ESPECIAL DO PETROLEO' in linha_upper: secao_atual = 'FEP'
            elif 'LC 176/2020' in linha_upper: secao_atual = 'LC_176'
            elif 'SIMPLES NACIONAL' in linha_upper: secao_atual = 'Simples_Nacional'
            elif 'PRODUCAO MINERAL' in linha_upper or 'CFM' in linha_upper: secao_atual = 'CFM'
            elif 'CIDE' in linha_upper: secao_atual = 'CIDE'
            
            if 'TOTAL POR PARCELA' in linha_upper:
                dentro_totais = True; continue
            if dentro_totais and ('DEBITO FUNDO' in linha_upper or 'CREDITO FUNDO' in linha_upper):
                dentro_totais = False
            if dentro_totais and 'RETENCAO PASEP' in linha_upper:
                match_valor = re.search(r'([\d._,]+)D', linha_upper)
                if match_valor and secao_atual in cliente_data:
                    cliente_data[secao_atual] = extrair_numero(match_valor.group(1))
        dados_lote.append(cliente_data)
    return pd.DataFrame(dados_lote)

# ── XML Injection Avançado (Bordas e Fundos Contábeis) ────────────────────────
def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_borders(cell, top=False, bottom=False, val="single", sz="6", color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    if top:
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), val)
        top_border.set(qn('w:sz'), sz)
        top_border.set(qn('w:space'), '0')
        top_border.set(qn('w:color'), color)
        tcBorders.append(top_border)

    if bottom:
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), val)
        bottom_border.set(qn('w:sz'), sz)
        bottom_border.set(qn('w:space'), '0')
        bottom_border.set(qn('w:color'), color)
        tcBorders.append(bottom_border)

# ==============================================================================
# MOTOR DE GERAÇÃO EXECUTIVA
# ==============================================================================
def gerar_documento_word_dinamico(template_bytes, dados):
    doc = Document(template_bytes)
    
    if dados.get('gerar_capa') and len(doc.paragraphs) > 0:
        p_base = doc.paragraphs[0]
        section = doc.sections[0]
        largura_util = section.page_width - section.left_margin - section.right_margin
        
        # 1. NÚMERO DO OFÍCIO E DATA
        p_cabecalho = p_base.insert_paragraph_before()
        p_cabecalho.paragraph_format.space_before = Pt(12)
        p_cabecalho.paragraph_format.space_after = Pt(24)
        
        tab_stops = p_cabecalho.paragraph_format.tab_stops
        tab_stops.add_tab_stop(largura_util, WD_TAB_ALIGNMENT.RIGHT)
        
        run_of = p_cabecalho.add_run(f"OFÍCIO Nº {str(dados['num_oficio']).strip()}")
        run_of.bold = True
        run_of.font.size = Pt(12)
        
        p_cabecalho.add_run("\t") 
        p_cabecalho.add_run(obter_data_formatada())
        
        # 2. DESTINATÁRIO BLINDADO (Com UF)
        p_destinatario = p_base.insert_paragraph_before()
        p_destinatario.paragraph_format.line_spacing = 1.0
        p_destinatario.paragraph_format.space_before = Pt(0)
        p_destinatario.paragraph_format.space_after = Pt(24) 
        
        p_destinatario.add_run("A Sua Excelência o(a) Senhor(a)\n")
        
        run_rep = p_destinatario.add_run(dados['representante_nome'] + "\n")
        run_rep.bold = True
        run_rep.font.size = Pt(12)
        
        p_destinatario.add_run("Representante Legal\n")
        p_destinatario.add_run(f"{dados['cliente']} - {dados['uf']}")

        # 3. ASSUNTO
        p_assunto = p_base.insert_paragraph_before()
        p_assunto.paragraph_format.space_after = Pt(24)
        run_ass_label = p_assunto.add_run("Assunto: ")
        run_ass_label.bold = True
        p_assunto.add_run(f"Demonstrativo Analítico de Apuração PASEP – Competência {dados['mes_ref']}.")
        
        # 4. VOCATIVO
        p_sauda = p_base.insert_paragraph_before()
        p_sauda.paragraph_format.first_line_indent = Cm(3.0)
        p_sauda.add_run("Senhor(a) Representante,")
        
        # 5. CORPO DO OFÍCIO
        corpo = (
            "\nCumprimentando-o(a) cordialmente, servimo-nos do presente para encaminhar o "
            "Demonstrativo Analítico de Apuração da Contribuição para o PASEP, bem como o "
            "detalhamento do cruzamento de dados de Malha Fiscal, referente à competência supramencionada.\n"
            "Ressaltamos que a análise leva em consideração as retenções na fonte efetuadas de forma automática pela "
            "Secretaria do Tesouro Nacional (STN), confrontando-as com as receitas correntes e as devidas exclusões legais.\n"
            "Colocamo-nos à inteira disposição para eventuais esclarecimentos e orientações técnicas adicionais.\n"
            "Respeitosamente,"
        )
        
        for paragrafo in corpo.strip().split('\n'):
            p_corpo = p_base.insert_paragraph_before(paragrafo.strip())
            p_corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_corpo.paragraph_format.first_line_indent = Cm(3.0) 
            p_corpo.paragraph_format.line_spacing = 1.5          
            p_corpo.paragraph_format.space_after = Pt(0)         
        
        # 6. ASSINATURA DUPLA
        p_espaco_ass = p_base.insert_paragraph_before()
        p_espaco_ass.paragraph_format.space_before = Pt(48)
        
        tabela_ass = doc.add_table(rows=1, cols=2)
        tabela_ass.style = 'Normal Table'
        
        p_espaco_ass._p.addnext(tabela_ass._tbl)
        
        c1 = tabela_ass.cell(0, 0)
        c2 = tabela_ass.cell(0, 1)
        
        metade_largura = largura_util / 2
        c1.width = metade_largura
        c2.width = metade_largura
        
        p_c1 = c1.paragraphs[0]
        p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_c1.add_run("Rubens Pires Malaquias\n").bold = True
        run_rubens_cargo = p_c1.add_run("Diretor Técnico e Consultor junto ao\nFisco Federal\nCRA/GO 6-007-48")
        run_rubens_cargo.font.size = Pt(10)
        
        p_c2 = c2.paragraphs[0]
        p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_c2.add_run("Glayzer Antônio Gomes da Silva\n").bold = True
        run_glayzer_cargo = p_c2.add_run("Advogado Especialista em Direito Público\nConstitucional e Administrativo\nOAB/GO 28.315")
        run_glayzer_cargo.font.size = Pt(10)
        
        p_break = p_base.insert_paragraph_before()
        p_break.add_run().add_break(WD_BREAK.PAGE)

    # --- SUBSTITUIÇÃO DE TAGS E TABELA ---
    tags_simples = {
        "{{ente}}": dados['cliente'],
        "{{competencia}}": dados['mes_ref'],
        "{{responsavel}}": dados.get('resp_apuracao', ''),
        "{{ramal}}": dados.get('ramal_apuracao', ''),
        "{{demanda}}": dados.get('num_demanda', '')
    }
    
    def substituir_no_bloco(bloco):
        txt = bloco.text
        for k, v in tags_simples.items():
            if k in txt:
                txt = txt.replace(k, str(v))
                bloco.text = txt
                
    for p in doc.paragraphs: substituir_no_bloco(p)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs: substituir_no_bloco(p)
                
    placeholder = "{{TABELA_APURACAO}}"
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = p.text.replace(placeholder, '')
            
            p_tab_titu = p.insert_paragraph_before()
            run_tab_titu = p_tab_titu.add_run(f"Memória Analítica de Cálculo – PASEP – {dados['mes_ref']}")
            run_tab_titu.bold = True
            p_tab_titu.add_run(f"\nMunicípio: {dados['cliente']} - {dados['uf']}")
            p_tab_titu.paragraph_format.space_after = Pt(18)
            
            tabela = doc.add_table(rows=0, cols=2)
            tabela.autofit = False
            
            # --- MOTOR DA TABELA CONTÁBIL (FLUXO DRE) ---
            def add_row(col1, col2, is_header=False, is_subtotal=False, is_final=False, is_indented=False, force_minus=False, obs=None, bg_color=None, is_risk=False):
                row = tabela.add_row().cells
                row[0].width = Cm(12.5) 
                row[1].width = Cm(3.5)
                
                p0 = row[0].paragraphs[0]
                p1 = row[1].paragraphs[0]
                
                p0.text = col1
                
                if obs:
                    run_obs = p0.add_run(f"\nNota: {obs}")
                    run_obs.italic = True
                    run_obs.font.size = Pt(9)
                    run_obs.font.color.rgb = RGBColor(120, 120, 120) 
                
                p1.text = col2
                
                p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                p0.paragraph_format.space_before = Pt(5)
                p0.paragraph_format.space_after = Pt(5)
                p1.paragraph_format.space_before = Pt(5)
                p1.paragraph_format.space_after = Pt(5)
                
                if is_header:
                    set_cell_background(row[0], "EAEAEA")
                    set_cell_background(row[1], "EAEAEA")
                    set_cell_borders(row[0], top=True, bottom=True, sz="8")
                    set_cell_borders(row[1], top=True, bottom=True, sz="8")
                    if p0.runs: p0.runs[0].bold = True
                    if p1.runs: p1.runs[0].bold = True
                    p0.paragraph_format.space_before = Pt(12)
                
                elif is_final:
                    set_cell_background(row[0], "F9F9F9")
                    set_cell_background(row[1], "F9F9F9")
                    set_cell_borders(row[0], top=True, bottom=True, val="double", sz="6")
                    set_cell_borders(row[1], top=True, bottom=True, val="double", sz="6")
                    if p0.runs: p0.runs[0].bold = True
                    if p1.runs: p1.runs[0].bold = True

                elif is_subtotal:
                    set_cell_borders(row[0], top=True, sz="6", color="808080")
                    set_cell_borders(row[1], top=True, sz="6", color="808080")
                    if p0.runs: p0.runs[0].bold = True
                    if p1.runs: p1.runs[0].bold = True
                
                if bg_color:
                    set_cell_background(row[0], bg_color)
                    set_cell_background(row[1], bg_color)

                if is_indented:
                    p0.paragraph_format.left_indent = Cm(0.5)
                
                # Regra estrita de cor vermelha (Apenas para deduções ou risco real aprovado)
                if (force_minus or is_risk) and p1.runs and p1.text != "R$ 0,00":
                    p1.runs[0].font.color.rgb = RGBColor(200, 0, 0)

            # ── FLUXO CONTÍNUO DRE ──
            add_row("Receitas Correntes e Capital", "Valores (R$)", is_header=True)
            add_row("1.0.00.00.00 - Receitas Correntes", formatar_moeda(dados['receitas_correntes']), is_indented=True)
            if dados['transferencias_capital'] > 0: add_row("2.4.00.00.00 - Transferências de Capital", formatar_moeda(dados['transferencias_capital']), is_indented=True)
            add_row("(I) Total da Receita", formatar_moeda(dados['total_receita']), is_subtotal=True)
            
            add_row("Deduções e Exclusões de Receitas", "", is_header=True)
            if dados['receitas_intra'] > 0: add_row("Dedução da Receita para formação do FUNDEB (91000)", formatar_moeda(dados['receitas_intra'], force_deduction_sign=True), is_indented=True, force_minus=True)
            for conv in dados['lista_convenios']:
                if conv['valor'] > 0: add_row(conv['desc'], formatar_moeda(conv['valor'], force_deduction_sign=True), is_indented=True, force_minus=True)
            add_row("(II) (-) Total das Deduções", formatar_moeda(dados['total_deducoes'], force_deduction_sign=True), is_subtotal=True)
            
            add_row("(III) BASE DE CÁLCULO BRUTA ( I - II )", formatar_moeda(dados['base_real']), is_final=True)
            
            add_row("A) PASEP Devido ( III * 1% )", formatar_moeda(dados['pasep_devido']), is_subtotal=True)
            
            add_row("B) (-) PASEP Retido na Fonte (STN)", formatar_moeda(dados['total_retido'], force_deduction_sign=True), is_subtotal=True, force_minus=True)
            if dados['ret_fpm'] > 0: add_row("FPM - Fundo de Participação dos Municípios", formatar_moeda(dados['ret_fpm']), is_indented=True)
            if dados['ret_fep'] > 0: add_row("FEP - Fundo Especial Petróleo", formatar_moeda(dados['ret_fep']), is_indented=True)
            if dados['ret_itr'] > 0: add_row("ITR - Imposto Territorial Rural", formatar_moeda(dados['ret_itr']), is_indented=True)
            if dados['ret_ado'] > 0: add_row("ADO - LC 176/2020", formatar_moeda(dados['ret_ado']), is_indented=True)
            if dados['ret_cid'] > 0: add_row("CIDE", formatar_moeda(dados['ret_cid']), is_indented=True)
            if dados['ret_cfm'] > 0: add_row("CFM (Mineral)", formatar_moeda(dados['ret_cfm']), is_indented=True)
            
            add_row("C) PASEP A RECOLHER (A - B)", formatar_moeda(dados['pasep_recolher']), is_final=True, bg_color="FCE4D6")
            
            add_row("Cruzamento Malha Fina (MIT x Apuração)", "", is_header=True)
            add_row("Valor Declarado no MIT", formatar_moeda(dados['valor_mit']), is_indented=True, obs=dados.get('obs_mit'))
            add_row("Valor Apurado (Tópico C)", formatar_moeda(dados['pasep_recolher']), is_indented=True)
            
            # Cálculo blindado do Risco
            diferenca_mit_c = round(dados['valor_mit'] - dados['pasep_recolher'], 2)
            risco = diferenca_mit_c != 0.00
            add_row("DIFERENÇA (Risco Fiscal)", formatar_moeda(diferenca_mit_c), is_final=True, is_risk=risco)
            
            # --- INJEÇÃO DA NOVA LINHA DE AUDITORIA (DENTRO DA TABELA) ---
            if dados.get('resp_apuracao') or dados.get('num_demanda'):
                row_controle = tabela.add_row().cells
                cell_controle = row_controle[0]
                cell_controle.merge(row_controle[1]) 
                
                p_c = cell_controle.paragraphs[0]
                p_c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_c.paragraph_format.space_before = Pt(6)
                p_c.paragraph_format.space_after = Pt(0)
                
                texto_controle = f"Apuração: {dados.get('resp_apuracao', '-')}  |  Ramal: {dados.get('ramal_apuracao', '-')}  |  Demanda: {dados.get('num_demanda', '-')}"
                run_c = p_c.add_run(texto_controle)
                run_c.italic = True
                run_c.font.size = Pt(8)
                run_c.font.color.rgb = RGBColor(128, 128, 128)
            
            p._p.addprevious(tabela._tbl)
            p._element.getparent().remove(p._element)
            break
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ── Interface Principal ────────────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Configurações Institucionais")
    
    modelo_word = st.file_uploader("1. Template Word (.docx)", type=["docx"])
    base_representantes = st.file_uploader("2. Lista Representantes (.csv)", type=["csv"])
    
    if base_representantes:
        try:
            st.session_state.df_representantes = pd.read_csv(base_representantes, sep=";")
            st.success("✅ Base de CNPJ carregada!")
        except Exception:
            st.error("Erro ao ler o CSV de Representantes.")

tab_import, tab_audit = st.tabs(["📂 1. Importação em Lote", "📊 2. Auditoria e Ofício CND"])

with tab_import:
    st.markdown("Arraste os arquivos CSV do Banco do Brasil. O processamento é **automático**.")
    arquivos_csv = st.file_uploader("Arquivos CSV", type=['csv'], accept_multiple_files=True, label_visibility="collapsed")
    
    if arquivos_csv:
        with st.spinner("Extraindo retenções detalhadas..."):
            st.session_state.df_clientes = processar_lote_csvs(arquivos_csv)
        st.success(f"✅ {len(arquivos_csv)} extrato(s) mapeado(s)!")

    if not st.session_state.df_clientes.empty:
        st.dataframe(st.session_state.df_clientes, use_container_width=True)

with tab_audit:
    col_input, col_preview = st.columns([1.2, 1])
    
    with col_input:
        st.header("📝 Dados do Cliente")
        
        lista_clientes = ["(Preenchimento Manual)"]
        if not st.session_state.df_clientes.empty:
            lista_clientes.extend(st.session_state.df_clientes['Cliente'].tolist())
            
        cliente_selecionado = st.selectbox("Selecione o Cliente", lista_clientes)
        
        dados_cliente = {}
        if cliente_selecionado != "(Preenchimento Manual)":
            dados_cliente = st.session_state.df_clientes[st.session_state.df_clientes['Cliente'] == cliente_selecionado].iloc[0].to_dict()
        
        col_ente, col_uf = st.columns([3, 1])
        with col_ente:
            cliente_nome = st.text_input("Ente", value=dados_cliente.get('Cliente', ''))
        with col_uf:
            cliente_uf = st.text_input("UF", value=dados_cliente.get('UF', 'GO'))
            
        mes_ref_padrao = dados_cliente.get('MesRef', 'Junho/2025')
        mes_ref = st.text_input("Competência", value=mes_ref_padrao)

        with st.expander("📋 Dados de Controle (Auditoria)", expanded=True):
            col_c1, col_c2, col_c3 = st.columns(3)
            with col_c1:
                resp_apuracao = st.text_input("Responsável (Analista)", placeholder="Ex: Samuel")
            with col_c2:
                ramal_apuracao = st.text_input("Ramal", placeholder="Ex: 1020")
            with col_c3:
                num_demanda = st.text_input("Nº da Demanda", value="20002000011")

        rep_encontrado = "Representante Legal"
        if not st.session_state.df_representantes.empty and cliente_nome:
            df_rep = st.session_state.df_representantes
            busca = df_rep[df_rep['Município'].str.contains(cliente_nome, case=False, na=False)]
            if not busca.empty:
                rep_encontrado = str(busca.iloc[0]['Responsável']).strip()

        with st.expander("✉️ Dados da Capa (Ofício Padrão RFB)", expanded=True):
            gerar_capa = st.checkbox("Adicionar Ofício Executivo CND no Word?", value=True)
            col_o1, col_o2 = st.columns(2)
            with col_o1:
                num_oficio = st.text_input("Número do Ofício", value="001/2026")
            with col_o2:
                representante_nome = st.text_input("Destinatário / Representante", value=rep_encontrado)

        with st.expander("1. Receitas e Deduções", expanded=True):
            rec_correntes = entrada_financeira("1.0.00.00.00 - Receitas Correntes", 0.0, key_suffix=f"rec_{cliente_selecionado}")
            transf_capital = entrada_financeira("2.4.00.00.00 - Transferências de Capital", 0.0, key_suffix=f"cap_{cliente_selecionado}")
            st.markdown("---")
            rec_intra = entrada_financeira("Dedução FUNDEB (91000)", 0.0, key_suffix=f"fun_{cliente_selecionado}")
            
            if st.button("➕ Adicionar Linha de Convênio"):
                st.session_state.num_convenios += 1

            lista_convenios_preenchidos = []
            for i in range(st.session_state.num_convenios):
                col_desc, col_val = st.columns([2, 1])
                with col_desc:
                    desc_conv = st.text_input("Descrição", value="Dedução de Convênios" if i == 0 else "", key=f"desc_conv_{cliente_selecionado}_{i}", label_visibility="collapsed" if i > 0 else "visible")
                with col_val:
                    val_conv = entrada_financeira("Valor (R$)", 0.0, key_suffix=f"val_conv_{cliente_selecionado}_{i}", label_visibility="collapsed" if i > 0 else "visible")
                if val_conv > 0:
                    lista_convenios_preenchidos.append({'desc': desc_conv.strip() or f"Convênio {i+1}", 'valor': val_conv})

        with st.expander("2. Retenções Analíticas", expanded=False):
            col_r1, col_r2 = st.columns(2)
            with col_r1:
                ret_fpm = entrada_financeira("FPM", dados_cliente.get('FPM', 0.0), key_suffix=cliente_selecionado)
                ret_fep = entrada_financeira("FEP", dados_cliente.get('FEP', 0.0), key_suffix=cliente_selecionado)
                ret_cid = entrada_financeira("CIDE", dados_cliente.get('CIDE', 0.0), key_suffix=cliente_selecionado)
            with col_r2:
                ret_itr = entrada_financeira("ITR", dados_cliente.get('ITR', 0.0), key_suffix=cliente_selecionado)
                ret_ado = entrada_financeira("ADO LC 176", dados_cliente.get('LC_176', 0.0), key_suffix=cliente_selecionado)
                ret_cfm = entrada_financeira("CFM (Mineral)", dados_cliente.get('CFM', 0.0), key_suffix=cliente_selecionado)

        with st.expander("3. MIT x Apuração (Tópico C)", expanded=False):
            valor_mit = entrada_financeira("Valor Informado no MIT", 0.0, key_suffix=cliente_selecionado)
            obs_mit = st.text_input("Nota Explicativa MIT (Opcional)", placeholder="Ex: Processo de Compensação Nº...", key=f"obs_mit_{cliente_selecionado}")

    # Lógica Tributária
    total_receita = rec_correntes + transf_capital
    total_convenios = sum(c['valor'] for c in lista_convenios_preenchidos)
    total_deducoes = rec_intra + total_convenios
    base_real = total_receita - total_deducoes
    pasep_devido = base_real * 0.01
    
    total_retido = ret_fpm + ret_fep + ret_cid + ret_itr + ret_ado + ret_cfm
    pasep_recolher = pasep_devido - total_retido if pasep_devido > total_retido else 0.0
    
    diferenca_mit_c = valor_mit - pasep_recolher

    with col_preview:
        st.header("📊 Prévia Fisco-Federal")
        
        st.markdown("### Composição da Base")
        st.text(f"(+) Receitas Totais: {formatar_moeda(total_receita)}\n(-) Deduções Totais: {formatar_moeda(total_deducoes)}\n(=) Base de Cálculo: {formatar_moeda(base_real)}")
        
        st.markdown("### Resumo Contínuo")
        st.info(f"**A) PASEP Devido (1%):** {formatar_moeda(pasep_devido)}\n**B) Retido na Fonte:** {formatar_moeda(total_retido)}\n**C) A Recolher:** {formatar_moeda(pasep_recolher)}")
        
        dados_doc = {
            'cliente': cliente_nome, 'uf': cliente_uf, 'mes_ref': mes_ref, 'receitas_correntes': rec_correntes,
            'transferencias_capital': transf_capital, 'total_receita': total_receita,
            'receitas_intra': rec_intra, 'lista_convenios': lista_convenios_preenchidos,
            'total_deducoes': total_deducoes, 'base_real': base_real,
            'ret_fpm': ret_fpm, 'ret_fep': ret_fep, 'ret_itr': ret_itr, 'ret_ado': ret_ado,
            'ret_cid': ret_cid, 'ret_cfm': ret_cfm, 'total_retido': total_retido,
            'pasep_devido': pasep_devido, 'pasep_recolher': pasep_recolher, 'valor_mit': valor_mit,
            'obs_mit': obs_mit, 'gerar_capa': gerar_capa, 'num_oficio': num_oficio, 
            'representante_nome': representante_nome,
            'resp_apuracao': resp_apuracao, 'ramal_apuracao': ramal_apuracao, 'num_demanda': num_demanda
        }

        if modelo_word:
            st.markdown("---")
            if st.button("🖨️ Gerar Ofício Executivo Completo", type="primary"):
                with st.spinner("Gerando Documento Ouro com Formatação DRE..."):
                    docx_pronto = gerar_documento_word_dinamico(modelo_word, dados_doc)
                st.success("✅ Ofício Ouro gerado com sucesso!")
                st.download_button(
                    label="📥 Baixar Documento Oficial (Word)",
                    data=docx_pronto,
                    file_name=f"Oficio_PASEP_{cliente_nome.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("⚠️ Insira o 'Folha_Rosto_PASEP_Modelo.docx' no menu lateral.")
